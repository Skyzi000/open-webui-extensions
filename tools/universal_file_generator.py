"""
title: Universal File Generator
author: Skyzi000 & Claude
version: 0.19.0
requirements: fastapi, python-docx, pandas, openpyxl, reportlab, weasyprint, beautifulsoup4, requests, markdown, pyzipper
description: |
  Universal file generation tool supporting unlimited text formats + binary formats with automatic cloud upload.
  
  ## Supported Formats
  - **Text**: All text-based formats (CSV, JSON, XML, TXT, HTML, Markdown, YAML, TOML, JavaScript, Python, SQL, etc.)
  - **Binary**: DOCX (with rich formatting), XLSX (Excel), PDF (with Japanese fonts), ZIP (with URL downloading and AES encryption support)
  - **Graphics**: SVG (scalable vector graphics with text, shapes, and structured elements)
  
  ## Key Features
  - Advanced PDF generation with WeasyPrint/ReportLab and Japanese font support
  - Rich DOCX formatting with HTML input support
  - SVG generation with structured elements (text, shapes, custom graphics)
  - ZIP archive creation with remote file downloading from URLs
  - Automatic cloud upload to multiple services (transfer.sh, 0x0.st, file.io, litterbox)
  - Comprehensive error handling and service fallback
  - BeautifulSoup HTML parsing for clean document generation
  - Pandas integration for advanced data processing
  
  ## Input Format Documentation
  Each file type expects specific data formats - see generate_file() docstring for detailed specifications.
"""

import json
import io
import zipfile
import requests
import base64
import mimetypes
import re
from typing import Awaitable, Callable, Dict, List, Any, Optional, Union
from datetime import datetime
from pydantic import BaseModel, Field
from fastapi import Request, UploadFile


# Optional dependencies with graceful fallback
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from bs4 import BeautifulSoup
    BS4_AVAILABLE = True
except ImportError:
    BS4_AVAILABLE = False

try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False

try:
    from weasyprint import HTML, CSS
    WEASYPRINT_AVAILABLE = True
except ImportError:
    WEASYPRINT_AVAILABLE = False

try:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib import colors
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.lib.fonts import addMapping
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    import markdown
    MARKDOWN_AVAILABLE = True
except ImportError:
    MARKDOWN_AVAILABLE = False

try:
    import pyzipper
    PYZIPPER_AVAILABLE = True
except ImportError:
    PYZIPPER_AVAILABLE = False

PDF_AVAILABLE = WEASYPRINT_AVAILABLE or REPORTLAB_AVAILABLE


class FileGenerator:
    """Internal file generation engine - not exposed to AI"""
    
    def __init__(self):
        pass

    def generate_content(self, file_type: str, data: Any, **kwargs) -> Optional[bytes]:
        """Generate file content based on type"""
        
        try:
            # Special binary formats that need custom processing
            if file_type == 'docx':
                return self.generate_docx(data, **kwargs)
            elif file_type == 'pdf':
                return self.generate_pdf(data, **kwargs)
            elif file_type == 'xlsx':
                return self.generate_xlsx(data, **kwargs)
            elif file_type == 'zip':
                return self.generate_zip(data, **kwargs)
            elif file_type == 'svg':
                return self.generate_svg(data, **kwargs)
            else:
                # Default: treat as text-based format
                return self.generate_text(data, **kwargs)
        except Exception as e:
            raise Exception(f"Content generation failed: {str(e)}")

    def get_mime_type(self, file_type: str) -> str:
        """Get MIME type for file format using Python standard library"""
        # Special handling for SVG to ensure correct MIME type
        if file_type.lower() == 'svg':
            return 'image/svg+xml'
        
        mime_type, _ = mimetypes.guess_type(f"dummy.{file_type}")
        return mime_type or 'text/plain'

    def generate_text(self, data: Union[str, Any], **kwargs) -> bytes:
        """Generate text content from string data"""
        if isinstance(data, str):
            # Convert escaped newlines to actual newlines
            text = data.replace('\\n', '\n').replace('\\t', '\t').replace('\\r', '\r')
            return text.encode('utf-8')
        else:
            # Fallback to string representation
            return str(data).encode('utf-8')









    def generate_docx(self, data: Union[str, Dict], **kwargs) -> bytes:
        """Generate DOCX content - accepts HTML, Markdown, or plain text"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required for DOCX generation. Install with: pip install python-docx")
        
        doc = Document()
        
        if isinstance(data, str):
            # Detect content type and convert to HTML if needed
            html_content = self._convert_to_html(data)
            if html_content:
                self._parse_html_to_docx(doc, html_content)
            else:
                self._add_text_content(doc, data)
        else:
            self._parse_any_structure(doc, data)
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()
    
    def _convert_to_html(self, content: str) -> str:
        """Convert Markdown or plain text to HTML, or return HTML if already HTML"""
        content = content.strip()
        
        # Already HTML
        if content.startswith('<') and '>' in content:
            return content
        
        # Detect Markdown patterns
        markdown_patterns = [
            r'^#+ ',  # Headers
            r'^\* ',  # Bullet lists
            r'^\d+\. ',  # Numbered lists
            r'\*\*.*?\*\*',  # Bold
            r'\*.*?\*',  # Italic
            r'`.*?`',  # Code
            r'^\|.*\|',  # Tables
            r'^\> ',  # Blockquotes
            r'```',  # Code blocks
        ]
        
        import re
        is_markdown = any(re.search(pattern, content, re.MULTILINE) for pattern in markdown_patterns)
        
        if is_markdown:
            return self._markdown_to_html(content)
        
        # Plain text - wrap in paragraphs
        lines = content.split('\n')
        html_lines = []
        for line in lines:
            line = line.strip()
            if line:
                html_lines.append(f'<p>{line}</p>')
            else:
                html_lines.append('<br>')
        
        return '\n'.join(html_lines)
    
    def _markdown_to_html(self, markdown_content: str) -> str:
        """Convert Markdown to HTML using markdown library"""
        if not MARKDOWN_AVAILABLE:
            raise ImportError("markdown library is required for Markdown conversion. Install with: pip install markdown")
        
        # Use proper markdown library with extensions
        return markdown.markdown(
            markdown_content,
            extensions=['tables', 'fenced_code', 'toc', 'codehilite']
        )
    
    def _parse_any_structure(self, doc, data):
        """Parse any dictionary structure intelligently"""
        if not isinstance(data, dict):
            doc.add_paragraph(str(data))
            return
        
        # Extract title first
        title_keys = ['title', 'タイトル', 'name', 'subject', 'heading']
        for key in title_keys:
            if key in data and data[key]:
                doc.add_heading(str(data[key]), 0)
                break
        
        # Extract metadata 
        meta_keys = ['date', 'author', 'created', 'updated', 'version']
        meta_info = []
        for key in meta_keys:
            if key in data and data[key]:
                meta_info.append(f"{key.title()}: {data[key]}")
        
        if meta_info:
            meta_p = doc.add_paragraph(" | ".join(meta_info))
            meta_p.italic = True
            doc.add_paragraph("")
        
        # Process remaining content
        processed = set(title_keys + meta_keys)
        
        for key, value in data.items():
            if key in processed:
                continue
                
            if isinstance(value, list):
                self._process_list(doc, key, value)
            elif isinstance(value, dict):
                doc.add_heading(str(key).replace('_', ' ').title(), 1)
                self._parse_any_structure(doc, value)
            elif value:
                self._add_key_value(doc, key, value)
    
    def _process_list(self, doc, key, items):
        """Process list items flexibly"""
        if not items:
            return
            
        # Don't add heading for common list names
        if key not in ['sections', 'content', 'items', 'data']:
            doc.add_heading(str(key).replace('_', ' ').title(), 1)
        
        for item in items:
            if isinstance(item, dict):
                self._process_dict_item(doc, item)
            elif isinstance(item, str) and item.strip():
                self._add_text_content(doc, item)
            elif item:
                doc.add_paragraph(str(item))
    
    def _process_dict_item(self, doc, item):
        """Process dictionary item looking for common patterns"""
        heading_keys = ['heading', 'title', 'name', 'section', 'header']
        content_keys = ['body', 'content', 'text', 'description', 'details']
        
        heading = None
        content = None
        
        # Find heading
        for h_key in heading_keys:
            if h_key in item and item[h_key]:
                heading = str(item[h_key])
                break
        
        # Find content  
        for c_key in content_keys:
            if c_key in item and item[c_key]:
                content = str(item[c_key])
                break
        
        if heading:
            doc.add_heading(heading, 2)
        
        if content:
            self._add_text_content(doc, content)
        
        # Process any remaining fields
        processed = set(heading_keys + content_keys)
        for key, value in item.items():
            if key not in processed and value:
                self._add_key_value(doc, key, value)
    
    def _add_key_value(self, doc, key, value):
        """Add key-value pair appropriately"""
        if isinstance(value, str) and len(value) > 50:
            doc.add_heading(str(key).replace('_', ' ').title(), 2)
            self._add_text_content(doc, value)
        else:
            doc.add_paragraph(f"{str(key).replace('_', ' ').title()}: {value}")
    
    def _add_text_content(self, doc, text):
        """Add text with proper line handling"""
        if not text:
            return
            
        # Handle various line break formats
        text = str(text).replace('\\n', '\n')
        lines = text.split('\n')
        
        for line in lines:
            line = line.strip()
            if line:
                doc.add_paragraph(line)

    def _parse_html_to_docx(self, doc, html_content):
        """Parse HTML to DOCX format using BeautifulSoup"""
        if not BS4_AVAILABLE:
            raise ImportError("BeautifulSoup4 is required for HTML parsing. Install with: pip install beautifulsoup4")
        
        # Parse HTML with BeautifulSoup
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Remove head and script tags
        for tag in soup(['head', 'script', 'style']):
            tag.decompose()
        
        # Find body content or use entire soup
        body = soup.find('body') or soup
        
        # Process each top-level element
        for element in body.descendants:
            if hasattr(element, 'name') and element.name and element.parent == body:
                self._process_html_element(doc, element)
    
    
    def _process_html_element(self, doc, element):
        """Process individual HTML elements with BeautifulSoup"""
        if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            # Handle headings
            level = int(element.name[1])
            text = element.get_text(strip=True)
            if text:
                doc.add_heading(text, level)
        
        elif element.name == 'p':
            # Handle paragraphs
            self._parse_html_element_to_docx(doc, element)
        
        elif element.name in ['ul', 'ol']:
            # Handle lists
            for li in element.find_all('li', recursive=False):
                self._parse_html_element_to_docx(doc, li, style='List Bullet')
        
        elif element.name == 'table':
            # Handle tables
            self._parse_html_table_to_docx_bs4(doc, element)
        
        elif element.name == 'div':
            # Handle divs - process children
            for child in element.children:
                if hasattr(child, 'name') and child.name:
                    self._process_html_element(doc, child)
                elif hasattr(child, 'strip') and child.strip():
                    doc.add_paragraph(child.strip())
        
        elif element.name == 'img':
            # Handle standalone images
            src = element.get('src', '')
            alt = element.get('alt', 'Image')
            if src:
                print(f"Found standalone image: src={src}, alt={alt}")
                # Create a paragraph for the image
                paragraph = doc.add_paragraph()
                self._add_image_to_paragraph(paragraph, src, alt)
        
        elif element.name == 'br':
            # Handle line breaks
            doc.add_paragraph()
        
        else:
            # Handle other elements as paragraphs
            text = element.get_text(strip=True)
            if text:
                doc.add_paragraph(text)

    def _parse_html_table_to_docx_bs4(self, doc, table_element):
        """Parse HTML table to DOCX format using BeautifulSoup"""
        rows = table_element.find_all('tr')
        if not rows:
            return
        
        # Extract table data with rich formatting
        table_data = []
        table_rich_data = []  # Store actual cell elements for rich formatting
        
        for row in rows:
            cells = row.find_all(['td', 'th'])
            if cells:
                cell_texts = []
                cell_elements = []
                for cell in cells:
                    # Check if cell contains links or other formatting
                    if cell.find('a') or cell.find('img') or cell.find(['strong', 'b', 'em', 'i']):
                        cell_elements.append(cell)
                        cell_texts.append('')  # Placeholder for rich content
                    else:
                        cell_elements.append(None)
                        cell_texts.append(cell.get_text(strip=True))
                table_data.append(cell_texts)
                table_rich_data.append(cell_elements)
        
        if table_data:
            # Add table to document
            table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
            table.style = 'Table Grid'
            
            for i, (row_data, rich_row) in enumerate(zip(table_data, table_rich_data)):
                row_cells = table.rows[i].cells
                for j, (cell_data, rich_cell) in enumerate(zip(row_data, rich_row)):
                    if j < len(row_cells):
                        if rich_cell:
                            # Handle rich content in cell
                            cell_paragraph = row_cells[j].paragraphs[0]
                            cell_paragraph.clear()  # Clear default content
                            self._parse_html_content_with_formatting_bs4(cell_paragraph, rich_cell)
                        else:
                            # Handle plain text
                            row_cells[j].text = cell_data
    

    def _parse_html_element_to_docx(self, doc, element, style=None):
        """Parse HTML element with links and images to DOCX using BeautifulSoup"""
        if not BS4_AVAILABLE:
            # Simple text stripping fallback
            import re
            text = re.sub(r'<[^>]+>', '', str(element)).strip()
            if text:
                if style:
                    paragraph = doc.add_paragraph(style=style)
                else:
                    paragraph = doc.add_paragraph()
                paragraph.add_run(text)
            return
        
        # Create a new paragraph
        if style:
            paragraph = doc.add_paragraph(style=style)
        else:
            paragraph = doc.add_paragraph()
        
        # Parse content with BeautifulSoup
        self._parse_html_content_with_formatting_bs4(paragraph, element)
    
    
    
    def _parse_html_content_with_formatting_bs4(self, paragraph, element):
        """Parse HTML content with BeautifulSoup and add formatting, links, and images"""
        # Simple recursive approach without complex tracking
        if hasattr(element, 'children'):
            for child in element.children:
                self._process_single_element_bs4(paragraph, child)
        else:
            # Handle case where element doesn't have children
            text = element.get_text() if hasattr(element, 'get_text') else str(element)
            if text.strip():
                paragraph.add_run(text.strip())
    
    def _process_single_element_bs4(self, paragraph, element):
        """Process a single element without complex nesting"""
        if hasattr(element, 'name') and element.name:
            if element.name == 'a':
                # Handle links
                href = element.get('href', '')
                text = element.get_text()
                if href and text:
                    self._add_hyperlink_to_paragraph(paragraph, href, text)
            elif element.name == 'img':
                # Handle images
                src = element.get('src', '')
                alt = element.get('alt', 'Image')
                if src:
                    self._add_image_to_paragraph(paragraph, src, alt)
            elif element.name == 'br':
                # Handle line breaks
                paragraph.add_run('\n')
            elif element.name in ['strong', 'b']:
                # Handle bold text
                text = element.get_text()
                if text.strip():
                    run = paragraph.add_run(text)
                    run.bold = True
            elif element.name in ['em', 'i']:
                # Handle italic text
                text = element.get_text()
                if text.strip():
                    run = paragraph.add_run(text)
                    run.italic = True
            else:
                # For other elements, recursively process children
                if hasattr(element, 'children'):
                    for child in element.children:
                        self._process_single_element_bs4(paragraph, child)
                else:
                    text = element.get_text()
                    if text.strip():
                        paragraph.add_run(text)
        elif hasattr(element, 'string') and element.string:
            # Handle text nodes
            text = element.string.strip()
            if text:
                paragraph.add_run(text)
        else:
            # Handle other text-like content
            text = str(element).strip()
            if text:
                paragraph.add_run(text)
    
    
    def _add_hyperlink_to_paragraph(self, paragraph, url, text):
        """Add hyperlink to paragraph"""
        try:
            # Create hyperlink relationship
            part = paragraph.part
            r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
            
            # Create hyperlink element
            hyperlink = paragraph._element.makeelement('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hyperlink')
            hyperlink.set('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id', r_id)
            
            # Create run element
            run = paragraph._element.makeelement('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r')
            run_props = paragraph._element.makeelement('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr')
            color = paragraph._element.makeelement('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}color')
            color.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', '0000FF')
            u = paragraph._element.makeelement('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}u')
            u.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 'single')
            run_props.append(color)
            run_props.append(u)
            run.append(run_props)
            
            text_elem = paragraph._element.makeelement('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
            text_elem.text = text
            run.append(text_elem)
            hyperlink.append(run)
            
            paragraph._element.append(hyperlink)
        except Exception:
            # Fallback: add as plain text with URL
            paragraph.add_run(f"{text} ({url})")
    
    def _add_image_to_paragraph(self, paragraph, src, alt_text):
        """Add image to paragraph"""
        try:
            # Try to download and add image
            if src.startswith(('http://', 'https://')):
                print(f"Downloading image from: {src}")
                response = requests.get(src, timeout=10, headers={'User-Agent': 'Mozilla/5.0'})
                if response.status_code == 200:
                    # Check if it's SVG content
                    content_type = response.headers.get('content-type', '').lower()
                    if 'svg' in content_type or src.lower().endswith('.svg'):
                        print(f"SVG images are not directly supported in DOCX: {alt_text}")
                        paragraph.add_run(f"[SVG Image: {alt_text}] (SVG not supported in DOCX - URL: {src})")
                        return
                        
                    image_stream = io.BytesIO(response.content)
                    try:
                        # Add image to paragraph with size constraints
                        run = paragraph.add_run()
                        from docx.shared import Inches
                        run.add_picture(image_stream, width=Inches(4))  # Set max width to 4 inches
                        print(f"Successfully added image: {alt_text}")
                        return
                    except Exception as e:
                        print(f"Failed to add image: {e}")
                        pass
                else:
                    print(f"Failed to download image: {response.status_code}")
            elif src.startswith('data:image/'):
                # Handle base64 images
                try:
                    print(f"Processing base64 image: {alt_text}")
                    header, data = src.split(',', 1)
                    # Check if it's SVG format
                    if 'svg' in header.lower():
                        print(f"SVG images are not directly supported in DOCX: {alt_text}")
                        paragraph.add_run(f"[SVG Image: {alt_text}] (SVG not supported in DOCX)")
                        return
                    
                    image_data = base64.b64decode(data)
                    image_stream = io.BytesIO(image_data)
                    run = paragraph.add_run()
                    from docx.shared import Inches
                    run.add_picture(image_stream, width=Inches(4))  # Set max width to 4 inches
                    print(f"Successfully added base64 image: {alt_text}")
                    return
                except Exception as e:
                    print(f"Failed to add base64 image: {e}")
                    pass
            else:
                print(f"Unsupported image source: {src}")
            
            # Fallback: add as text
            paragraph.add_run(f"[Image: {alt_text}]")
            print(f"Added image as text fallback: {alt_text}")
        except Exception as e:
            print(f"Exception in _add_image_to_paragraph: {e}")
            # Fallback: add as text
            paragraph.add_run(f"[Image: {alt_text}]")

    def generate_pdf(self, data: Union[str, Dict, List], **kwargs) -> bytes:
        """Generate PDF content with flexible data structure support"""
        if not PDF_AVAILABLE:
            raise ImportError("PDF generation requires weasyprint or reportlab. Install with: pip install weasyprint or pip install reportlab")
        
        # Prefer WeasyPrint for better HTML/CSS support
        if WEASYPRINT_AVAILABLE:
            return self._generate_pdf_with_weasyprint(data)
        elif REPORTLAB_AVAILABLE:
            return self._generate_pdf_with_html(data)
        else:
            raise ImportError("No PDF generation library available")
    
    def _generate_pdf_with_weasyprint(self, data) -> bytes:
        """Generate PDF using WeasyPrint for superior HTML/CSS rendering"""
        # Ensure Japanese fonts are available
        font_css = self._ensure_japanese_fonts_for_weasyprint()
        
        # Convert data to HTML
        if isinstance(data, str):
            # Detect content type and convert to HTML if needed
            html_content = self._convert_to_html(data)
        elif isinstance(data, dict):
            html_content = self._dict_to_html_full(data)
        elif isinstance(data, list):
            html_content = self._list_to_html_full(data)
        else:
            html_content = f"<p>{str(data)}</p>"
        
        # Create complete HTML document with enhanced CSS
        full_html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <style>
                {font_css}
                
                @page {{
                    size: A4;
                    margin: 2cm;
                }}
                
                body {{
                    font-family: 'NotoSansJP', 'DejaVu Sans', sans-serif;
                    font-size: 11px;
                    line-height: 1.6;
                    color: #333;
                    margin: 0;
                    padding: 0;
                }}
                
                h1 {{
                    font-size: 20px;
                    color: #2c3e50;
                    border-bottom: 3px solid #3498db;
                    padding-bottom: 10px;
                    margin-bottom: 25px;
                    page-break-after: avoid;
                }}
                
                h2 {{
                    font-size: 16px;
                    color: #34495e;
                    margin-top: 30px;
                    margin-bottom: 15px;
                    page-break-after: avoid;
                }}
                
                h3 {{
                    font-size: 14px;
                    color: #7f8c8d;
                    margin-top: 20px;
                    margin-bottom: 10px;
                    page-break-after: avoid;
                }}
                
                p {{
                    margin-bottom: 12px;
                    text-align: justify;
                }}
                
                ul, ol {{
                    margin-bottom: 15px;
                    padding-left: 25px;
                }}
                
                li {{
                    margin-bottom: 6px;
                }}
                
                table {{
                    width: 100%;
                    border-collapse: collapse;
                    margin: 20px 0;
                    page-break-inside: avoid;
                    box-shadow: 0 2px 8px rgba(0,0,0,0.1);
                }}
                
                th {{
                    background-color: #34495e;
                    color: white;
                    padding: 12px;
                    text-align: left;
                    font-weight: bold;
                    border: 1px solid #2c3e50;
                }}
                
                td {{
                    padding: 10px 12px;
                    border: 1px solid #ecf0f1;
                    vertical-align: top;
                }}
                
                tr:nth-child(even) {{
                    background-color: #f8f9fa;
                }}
                
                strong {{
                    color: #2c3e50;
                    font-weight: bold;
                }}
                
                .header {{
                    text-align: center;
                    margin-bottom: 40px;
                }}
                
                .page-break {{
                    page-break-before: always;
                }}
                
                .no-break {{
                    page-break-inside: avoid;
                }}
                
                code {{
                    background-color: #f4f4f4;
                    padding: 2px 4px;
                    border-radius: 3px;
                    font-family: 'Courier New', monospace;
                    font-size: 10px;
                }}
                
                blockquote {{
                    border-left: 4px solid #3498db;
                    margin: 15px 0;
                    padding-left: 15px;
                    color: #555;
                    font-style: italic;
                }}
            </style>
        </head>
        <body>
            {html_content}
        </body>
        </html>
        """
        
        try:
            # Generate PDF using WeasyPrint with font subsetting disabled
            html_doc = HTML(string=full_html)
            
            # Try with font configuration for newer versions
            try:
                from weasyprint.fonts import FontConfiguration
                font_config = FontConfiguration()
                # Disable font subsetting to preserve all Japanese glyphs
                pdf_bytes = html_doc.write_pdf(
                    font_config=font_config, 
                    optimize_images=False,
                    presentational_hints=True
                )
            except ImportError:
                # Fallback for older WeasyPrint versions
                print("Using older WeasyPrint version without FontConfiguration")
                pdf_bytes = html_doc.write_pdf()
            
            return pdf_bytes
        except Exception as e:
            print(f"WeasyPrint error: {e}")
            # Fallback to reportlab if available
            if REPORTLAB_AVAILABLE:
                print("Falling back to reportlab...")
                return self._generate_pdf_with_html(data)
            else:
                raise e
    
    def _ensure_japanese_fonts_for_weasyprint(self) -> str:
        """Setup Japanese fonts for WeasyPrint - try system fonts first, then download"""
        import tempfile
        import os
        
        # Try system fonts first (common in Docker containers)
        system_fonts = [
            "Noto Sans CJK JP",
            "Noto Sans JP", 
            "DejaVu Sans",
            "Liberation Sans"
        ]
        
        font_css = f"""
            body, h1, h2, h3, h4, h5, h6, p, td, th, li {{
                font-family: {', '.join([f"'{font}'" for font in system_fonts])}, sans-serif !important;
            }}
        """
        
        # Try to download a comprehensive Japanese font
        temp_dir = tempfile.gettempdir()
        font_path = os.path.join(temp_dir, 'NotoSansCJK-Regular.otf')
        
        # Check if font already exists
        if os.path.exists(font_path):
            print(f"✓ Japanese font already available: {font_path}")
            font_css = f"""
                @font-face {{
                    font-family: 'NotoSansCJK';
                    src: url('file://{font_path}') format('opentype');
                    font-weight: normal;
                    font-style: normal;
                }}
                
                body, h1, h2, h3, h4, h5, h6, p, td, th, li {{
                    font-family: 'NotoSansCJK', 'DejaVu Sans', sans-serif !important;
                    font-feature-settings: normal;
                }}
            """
            return font_css
        
        print("=== Downloading comprehensive Japanese font for WeasyPrint ===")
        
        # Try to download comprehensive CJK font - use smaller files for Docker
        font_urls = [
            "https://github.com/notofonts/noto-cjk/raw/main/Sans/OTF/Japanese/NotoSansCJKjp-Regular.otf",
            "https://fonts.gstatic.com/ea/notosansjapanese/v6/NotoSansJP-Regular.otf"
        ]
        
        for i, font_url in enumerate(font_urls):
            try:
                print(f"Trying comprehensive font source {i+1}: {font_url}")
                response = requests.get(font_url, timeout=30)
                if response.status_code == 200:
                    with open(font_path, 'wb') as f:
                        f.write(response.content)
                    
                    print(f"✓ Comprehensive Japanese font downloaded: {font_path} ({len(response.content)} bytes)")
                    
                    # Create CSS with comprehensive font-face declaration
                    font_css = f"""
                        @font-face {{
                            font-family: 'NotoSansCJK';
                            src: url('file://{font_path}') format('opentype');
                            font-weight: normal;
                            font-style: normal;
                        }}
                        
                        body, h1, h2, h3, h4, h5, h6, p, td, th, li {{
                            font-family: 'NotoSansCJK', 'DejaVu Sans', sans-serif !important;
                            font-feature-settings: normal;
                        }}
                    """
                    return font_css
                    
            except Exception as e:
                print(f"Comprehensive font source {i+1} failed: {e}")
                continue
        
        print("✗ Could not download comprehensive Japanese font, using system fonts")
        return font_css
    
    def _dict_to_html_full(self, data):
        """Convert dictionary to complete HTML with better formatting"""
        html_parts = []
        
        # Look for title
        title_keys = ['title', 'タイトル', 'name', 'subject', 'heading']
        for key in title_keys:
            if key in data and data[key]:
                html_parts.append(f'<div class="header"><h1>{data[key]}</h1></div>')
                break
        
        # Process other content
        processed_keys = set(title_keys)
        
        for key, value in data.items():
            if key in processed_keys:
                continue
            
            if isinstance(value, list):
                html_parts.append(f"<h2>{key.replace('_', ' ').title()}</h2>")
                
                # Check if it's table data
                if value and isinstance(value[0], dict) and len(value) > 1:
                    # Table format
                    html_parts.append(self._create_html_table(value))
                else:
                    # List format
                    html_parts.append("<ul>")
                    for item in value:
                        if isinstance(item, dict):
                            html_parts.append(f"<li>{self._dict_to_html_inline(item)}</li>")
                        else:
                            html_parts.append(f"<li>{item}</li>")
                    html_parts.append("</ul>")
            
            elif isinstance(value, dict):
                html_parts.append(f"<h2>{key.replace('_', ' ').title()}</h2>")
                html_parts.append(self._dict_to_html_full(value))
            
            else:
                if len(str(value)) > 50:
                    html_parts.append(f"<h3>{key.replace('_', ' ').title()}</h3>")
                    html_parts.append(f"<p>{value}</p>")
                else:
                    html_parts.append(f"<p><strong>{key.replace('_', ' ').title()}:</strong> {value}</p>")
        
        return '\n'.join(html_parts)
    
    def _list_to_html_full(self, data):
        """Convert list to complete HTML"""
        html_parts = []
        
        # Check if it's table data
        if data and isinstance(data[0], dict) and len(data) > 1:
            # Table format
            html_parts.append(self._create_html_table(data))
        else:
            # List format
            html_parts.append("<ul>")
            for item in data:
                if isinstance(item, dict):
                    html_parts.append(f"<li>{self._dict_to_html_inline(item)}</li>")
                else:
                    html_parts.append(f"<li>{item}</li>")
            html_parts.append("</ul>")
        
        return '\n'.join(html_parts)
    
    def _dict_to_html_inline(self, data):
        """Convert dictionary to inline HTML"""
        parts = []
        for key, value in data.items():
            if isinstance(value, (dict, list)):
                continue  # Skip complex structures in inline mode
            parts.append(f"<strong>{key.replace('_', ' ').title()}:</strong> {value}")
        return " | ".join(parts)
    
    def _create_html_table(self, data):
        """Create HTML table from list of dictionaries"""
        if not data:
            return ""
        
        # Get headers from first item
        headers = list(data[0].keys())
        
        html = ["<table>"]
        
        # Add header row
        html.append("<tr>")
        for header in headers:
            html.append(f"<th>{header.replace('_', ' ').title()}</th>")
        html.append("</tr>")
        
        # Add data rows
        for row in data:
            html.append("<tr>")
            for header in headers:
                value = row.get(header, "")
                html.append(f"<td>{value}</td>")
            html.append("</tr>")
        
        html.append("</table>")
        
        return '\n'.join(html)

    def _generate_pdf_with_html(self, data) -> bytes:
        """Generate PDF using Platypus with HTML parsing"""
        buffer = io.BytesIO()
        
        # Force download and register Japanese font
        print("=== PDF Generation with HTML and Japanese Font Support ===")
        font_registered = self._download_and_register_japanese_font()
        self._font_registered = font_registered
        
        # Create document with Platypus
        doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=50, bottomMargin=50, 
                               leftMargin=50, rightMargin=50)
        
        # Get base styles
        styles = getSampleStyleSheet()
        
        # Create custom styles with Japanese font support
        if font_registered:
            japanese_font = 'NotoSansJP'
            print(f"✓ Using Japanese font: {japanese_font}")
        else:
            japanese_font = 'Helvetica'
            print(f"✗ Japanese font failed, using: {japanese_font}")
        
        # Create custom styles
        custom_styles = {
            'Title': ParagraphStyle(
                'CustomTitle',
                parent=styles['Title'],
                fontName=japanese_font,
                fontSize=18,
                spaceAfter=20,
                alignment=1,  # Center
            ),
            'Heading1': ParagraphStyle(
                'CustomHeading1',
                parent=styles['Heading1'],
                fontName=japanese_font,
                fontSize=14,
                spaceAfter=12,
                spaceBefore=12,
            ),
            'Heading2': ParagraphStyle(
                'CustomHeading2',
                parent=styles['Heading2'],
                fontName=japanese_font,
                fontSize=12,
                spaceAfter=10,
                spaceBefore=10,
            ),
            'Normal': ParagraphStyle(
                'CustomNormal',
                parent=styles['Normal'],
                fontName=japanese_font,
                fontSize=10,
                spaceAfter=6,
                wordWrap='CJK',
            ),
            'Bullet': ParagraphStyle(
                'CustomBullet',
                parent=styles['Normal'],
                fontName=japanese_font,
                fontSize=10,
                spaceAfter=6,
                leftIndent=20,
                bulletIndent=10,
                wordWrap='CJK',
            ),
        }
        
        # Create story (list of flowables)
        story = []
        
        # Process data to create flowables
        processed_data = self._process_html_data(data, custom_styles)
        story.extend(processed_data)
        
        # Build the PDF
        doc.build(story)
        
        # Get the PDF content
        buffer.seek(0)
        return buffer.read()
    
    def _process_html_data(self, data, styles):
        """Process HTML data into Platypus flowables"""
        flowables = []
        
        if isinstance(data, str):
            # Detect content type and convert to HTML if needed
            html_content = self._convert_to_html(data)
            flowables.extend(self._parse_html_to_flowables(html_content, styles))
        
        elif isinstance(data, dict):
            # Handle dictionary data - convert to HTML first
            html_content = self._dict_to_html(data)
            flowables.extend(self._parse_html_to_flowables(html_content, styles))
        
        elif isinstance(data, list):
            # Handle list data - convert to HTML first
            html_content = self._list_to_html(data)
            flowables.extend(self._parse_html_to_flowables(html_content, styles))
        
        return flowables
    
    def _parse_html_to_flowables(self, html_text, styles):
        """Parse HTML text into Platypus flowables"""
        flowables = []
        
        # Simple HTML parsing - handle basic tags
        import re
        
        # Replace common HTML tags with Platypus Paragraph styles
        html_text = str(html_text)
        
        # Split by major block elements
        blocks = re.split(r'(<h[1-6]>.*?</h[1-6]>|<p>.*?</p>|<ul>.*?</ul>|<ol>.*?</ol>|<table>.*?</table>)', html_text, flags=re.DOTALL)
        
        for block in blocks:
            block = block.strip()
            if not block:
                continue
            
            # Handle headings
            if re.match(r'<h1>.*?</h1>', block, re.DOTALL):
                text = re.sub(r'<h1>(.*?)</h1>', r'\1', block, flags=re.DOTALL)
                flowables.append(Paragraph(text, styles['Title']))
            elif re.match(r'<h[2-3]>.*?</h[2-3]>', block, re.DOTALL):
                text = re.sub(r'<h[2-3]>(.*?)</h[2-3]>', r'\1', block, flags=re.DOTALL)
                flowables.append(Paragraph(text, styles['Heading1']))
            elif re.match(r'<h[4-6]>.*?</h[4-6]>', block, re.DOTALL):
                text = re.sub(r'<h[4-6]>(.*?)</h[4-6]>', r'\1', block, flags=re.DOTALL)
                flowables.append(Paragraph(text, styles['Heading2']))
            
            # Handle paragraphs
            elif re.match(r'<p>.*?</p>', block, re.DOTALL):
                text = re.sub(r'<p>(.*?)</p>', r'\1', block, flags=re.DOTALL)
                flowables.append(Paragraph(text, styles['Normal']))
            
            # Handle lists
            elif re.match(r'<ul>.*?</ul>', block, re.DOTALL):
                list_items = re.findall(r'<li>(.*?)</li>', block, re.DOTALL)
                for item in list_items:
                    flowables.append(Paragraph(f"• {item.strip()}", styles['Bullet']))
            
            elif re.match(r'<ol>.*?</ol>', block, re.DOTALL):
                list_items = re.findall(r'<li>(.*?)</li>', block, re.DOTALL)
                for i, item in enumerate(list_items, 1):
                    flowables.append(Paragraph(f"{i}. {item.strip()}", styles['Bullet']))
            
            # Handle tables
            elif re.match(r'<table>.*?</table>', block, re.DOTALL):
                table_flowable = self._parse_html_table(block)
                if table_flowable:
                    flowables.append(table_flowable)
                    flowables.append(Spacer(1, 12))
            
            # Handle plain text
            else:
                if block.strip():
                    flowables.append(Paragraph(block, styles['Normal']))
        
        return flowables
    
    def _parse_html_table(self, table_html):
        """Parse HTML table into Table flowable"""
        try:
            import re
            
            # Extract table rows
            rows = re.findall(r'<tr>(.*?)</tr>', table_html, re.DOTALL)
            table_data = []
            
            for row in rows:
                # Extract cells (th or td)
                cells = re.findall(r'<(?:th|td)>(.*?)</(?:th|td)>', row, re.DOTALL)
                if cells:
                    table_data.append([cell.strip() for cell in cells])
            
            if table_data:
                return self._create_table_flowable_from_data(table_data)
            
        except Exception as e:
            print(f"Error parsing HTML table: {e}")
            return None
    
    def _dict_to_html(self, data):
        """Convert dictionary to HTML text"""
        html_parts = []
        
        # Look for title
        title_keys = ['title', 'タイトル', 'name', 'subject', 'heading']
        for key in title_keys:
            if key in data and data[key]:
                html_parts.append(f"<h1>{data[key]}</h1>")
                break
        
        # Process other content
        processed_keys = set(title_keys)
        
        for key, value in data.items():
            if key in processed_keys:
                continue
            
            if isinstance(value, list):
                html_parts.append(f"<h2>{key.replace('_', ' ').title()}</h2>")
                html_parts.append("<ul>")
                for item in value:
                    if isinstance(item, dict):
                        html_parts.append(f"<li>{self._dict_to_html(item)}</li>")
                    else:
                        html_parts.append(f"<li>{item}</li>")
                html_parts.append("</ul>")
            
            elif isinstance(value, dict):
                html_parts.append(f"<h2>{key.replace('_', ' ').title()}</h2>")
                html_parts.append(self._dict_to_html(value))
            
            else:
                if len(str(value)) > 50:
                    html_parts.append(f"<h2>{key.replace('_', ' ').title()}</h2>")
                    html_parts.append(f"<p>{value}</p>")
                else:
                    html_parts.append(f"<p><strong>{key.replace('_', ' ').title()}:</strong> {value}</p>")
        
        return '\n'.join(html_parts)
    
    def _list_to_html(self, data):
        """Convert list to HTML text"""
        html_parts = []
        
        html_parts.append("<ul>")
        for item in data:
            if isinstance(item, dict):
                html_parts.append(f"<li>{self._dict_to_html(item)}</li>")
            else:
                html_parts.append(f"<li>{item}</li>")
        html_parts.append("</ul>")
        
        return '\n'.join(html_parts)

    def _generate_pdf_with_platypus_old(self, data) -> bytes:
        """Generate PDF using Platypus for better document flow"""
        buffer = io.BytesIO()
        
        # Force download and register Japanese font with detailed logging
        print("=== PDF Generation with Platypus and Japanese Font Support ===")
        font_registered = self._download_and_register_japanese_font()
        self._font_registered = font_registered
        
        # Create document with Platypus
        doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=50, bottomMargin=50, 
                               leftMargin=50, rightMargin=50)
        
        # Get base styles
        styles = getSampleStyleSheet()
        
        # Create custom styles with Japanese font support
        if font_registered:
            japanese_font = 'NotoSansJP'
            print(f"✓ Using Japanese font: {japanese_font}")
        else:
            japanese_font = 'Helvetica'
            print(f"✗ Japanese font failed, using: {japanese_font}")
        
        # Create custom styles
        custom_styles = {
            'Title': ParagraphStyle(
                'CustomTitle',
                parent=styles['Title'],
                fontName=japanese_font,
                fontSize=18,
                spaceAfter=20,
                alignment=1,  # Center
            ),
            'Heading1': ParagraphStyle(
                'CustomHeading1',
                parent=styles['Heading1'],
                fontName=japanese_font,
                fontSize=14,
                spaceAfter=12,
                spaceBefore=12,
            ),
            'Heading2': ParagraphStyle(
                'CustomHeading2',
                parent=styles['Heading2'],
                fontName=japanese_font,
                fontSize=12,
                spaceAfter=10,
                spaceBefore=10,
            ),
            'Normal': ParagraphStyle(
                'CustomNormal',
                parent=styles['Normal'],
                fontName=japanese_font,
                fontSize=10,
                spaceAfter=6,
                wordWrap='CJK',
            ),
            'Bullet': ParagraphStyle(
                'CustomBullet',
                parent=styles['Normal'],
                fontName=japanese_font,
                fontSize=10,
                spaceAfter=6,
                leftIndent=20,
                bulletIndent=10,
                wordWrap='CJK',
            ),
        }
        
        # Create story (list of flowables)
        story = []
        
        # Process data to create flowables
        processed_data = self._process_data_for_platypus(data, custom_styles)
        story.extend(processed_data)
        
        # Build the PDF
        doc.build(story)
        
        # Get the PDF content
        buffer.seek(0)
        return buffer.read()
    
    def _process_data_for_platypus(self, data, styles):
        """Process data into Platypus flowables - focused on Markdown parsing"""
        flowables = []
        
        if isinstance(data, str):
            # Handle Markdown text
            flowables.extend(self._parse_markdown_to_flowables(data, styles))
        
        elif isinstance(data, dict):
            # Handle dictionary data - convert to text first
            text_content = self._dict_to_markdown(data)
            flowables.extend(self._parse_markdown_to_flowables(text_content, styles))
        
        elif isinstance(data, list):
            # Handle list data - convert to text first
            text_content = self._list_to_markdown(data)
            flowables.extend(self._parse_markdown_to_flowables(text_content, styles))
        
        return flowables
    
    def _parse_markdown_to_flowables(self, markdown_text, styles):
        """Parse Markdown text into Platypus flowables"""
        flowables = []
        lines = markdown_text.split('\n')
        
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            
            if not line:
                # Empty line - add small spacer
                flowables.append(Spacer(1, 6))
                i += 1
                continue
            
            # Check for table (starts with |)
            if line.startswith('|') and '|' in line[1:]:
                # Parse table
                table_lines = []
                j = i
                while j < len(lines) and lines[j].strip().startswith('|'):
                    table_lines.append(lines[j].strip())
                    j += 1
                
                if len(table_lines) > 1:
                    # Find separator line (usually second line with |---|)
                    separator_index = -1
                    for idx, line in enumerate(table_lines):
                        if '---' in line or '--' in line:
                            separator_index = idx
                            break
                    
                    if separator_index > 0:
                        # Process table with separator
                        headers = [cell.strip() for cell in table_lines[0].split('|')[1:-1]]
                        data_rows = []
                        for row_line in table_lines[separator_index + 1:]:
                            row_data = [cell.strip() for cell in row_line.split('|')[1:-1]]
                            data_rows.append(row_data)
                        
                        table_data = [headers] + data_rows
                        table_flowable = self._create_table_flowable_from_data(table_data)
                        if table_flowable:
                            flowables.append(table_flowable)
                            flowables.append(Spacer(1, 12))
                    else:
                        # Process table without separator (all rows are data)
                        table_data = []
                        for row_line in table_lines:
                            row_data = [cell.strip() for cell in row_line.split('|')[1:-1]]
                            table_data.append(row_data)
                        
                        table_flowable = self._create_table_flowable_from_data(table_data)
                        if table_flowable:
                            flowables.append(table_flowable)
                            flowables.append(Spacer(1, 12))
                    
                    i = j
                    continue
            
            # Handle headings
            if line.startswith('# '):
                flowables.append(Paragraph(line[2:], styles['Title']))
            elif line.startswith('## '):
                flowables.append(Paragraph(line[3:], styles['Heading1']))
            elif line.startswith('### '):
                flowables.append(Paragraph(line[4:], styles['Heading2']))
            elif line.startswith('- ') or line.startswith('* '):
                # Bullet point
                flowables.append(Paragraph(line[2:], styles['Bullet']))
            else:
                # Normal text
                flowables.append(Paragraph(line, styles['Normal']))
            
            i += 1
        
        return flowables
    
    def _create_table_flowable_from_data(self, table_data):
        """Create a Table flowable from parsed table data"""
        try:
            # Create table with style (use Japanese font for table content)
            # Make a copy to avoid modifying the original data
            table_data_copy = [row[:] for row in table_data]
            japanese_font = 'NotoSansJP' if self._font_registered else 'Helvetica'
            
            # Use HTML bold tags for Japanese font headers
            if self._font_registered and table_data_copy:
                # For Japanese font, use HTML bold tags in the header text
                for i, header in enumerate(table_data_copy[0]):
                    table_data_copy[0][i] = f"<b>{header}</b>"
            
            table = Table(table_data_copy)
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), f'{japanese_font}-Bold' if japanese_font == 'Helvetica' else f'{japanese_font}'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('FONTNAME', (0, 1), (-1, -1), japanese_font),
                ('FONTSIZE', (0, 1), (-1, -1), 10),
                ('ALIGN', (0, 1), (-1, -1), 'LEFT'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.lightgrey]),
            ]))
            
            return table
            
        except Exception as e:
            print(f"Error creating table: {e}")
            return None
    
    def _dict_to_markdown(self, data):
        """Convert dictionary to Markdown text"""
        lines = []
        
        # Look for title
        title_keys = ['title', 'タイトル', 'name', 'subject', 'heading']
        for key in title_keys:
            if key in data and data[key]:
                lines.append(f"# {data[key]}")
                lines.append("")
                break
        
        # Process other content
        processed_keys = set(title_keys)
        
        for key, value in data.items():
            if key in processed_keys:
                continue
            
            if isinstance(value, list):
                lines.append(f"## {key.replace('_', ' ').title()}")
                lines.append("")
                for item in value:
                    if isinstance(item, dict):
                        lines.append(self._dict_to_markdown(item))
                    else:
                        lines.append(f"- {item}")
                lines.append("")
            
            elif isinstance(value, dict):
                lines.append(f"## {key.replace('_', ' ').title()}")
                lines.append("")
                lines.append(self._dict_to_markdown(value))
                lines.append("")
            
            else:
                if len(str(value)) > 50:
                    lines.append(f"## {key.replace('_', ' ').title()}")
                    lines.append("")
                    lines.append(str(value))
                else:
                    lines.append(f"**{key.replace('_', ' ').title()}:** {value}")
                lines.append("")
        
        return '\n'.join(lines)
    
    def _list_to_markdown(self, data):
        """Convert list to Markdown text"""
        lines = []
        
        for item in data:
            if isinstance(item, dict):
                lines.append(self._dict_to_markdown(item))
                lines.append("")
            else:
                lines.append(f"- {item}")
        
        return '\n'.join(lines)
    
    def _process_dict_for_platypus(self, data, flowables, styles):
        """Process dictionary data into Platypus flowables"""
        
        # Look for title/heading
        title_keys = ['title', 'タイトル', 'name', 'subject', 'heading']
        for key in title_keys:
            if key in data and data[key]:
                flowables.append(Paragraph(str(data[key]), styles['Title']))
                flowables.append(Spacer(1, 12))
                break
        
        # Look for tables
        if 'table' in data or 'tables' in data:
            table_data = data.get('table') or data.get('tables')
            if table_data:
                table_flowable = self._create_table_flowable(table_data, styles)
                if table_flowable:
                    flowables.append(table_flowable)
                    flowables.append(Spacer(1, 12))
        
        # Process other content
        processed_keys = set(title_keys + ['table', 'tables'])
        
        for key, value in data.items():
            if key in processed_keys:
                continue
            
            if isinstance(value, list):
                # Section heading
                flowables.append(Paragraph(str(key).replace('_', ' ').title(), styles['Heading1']))
                
                # Process list items
                for item in value:
                    if isinstance(item, dict):
                        self._process_dict_for_platypus(item, flowables, styles)
                        flowables.append(Spacer(1, 6))
                    else:
                        flowables.append(Paragraph(f"• {str(item)}", styles['Bullet']))
                
                flowables.append(Spacer(1, 12))
            
            elif isinstance(value, dict):
                # Subsection
                flowables.append(Paragraph(str(key).replace('_', ' ').title(), styles['Heading1']))
                self._process_dict_for_platypus(value, flowables, styles)
                flowables.append(Spacer(1, 12))
            
            else:
                # Key-value pair
                if len(str(value)) > 50:
                    # Long value - use as paragraph with key as heading
                    flowables.append(Paragraph(str(key).replace('_', ' ').title(), styles['Heading2']))
                    flowables.append(Paragraph(str(value), styles['Normal']))
                else:
                    # Short value - inline
                    flowables.append(Paragraph(f"<b>{str(key).replace('_', ' ').title()}:</b> {str(value)}", styles['Normal']))
                
                flowables.append(Spacer(1, 6))
    
    def _create_table_flowable(self, table_data, styles):
        """Create a Table flowable from table data"""
        try:
            if isinstance(table_data, dict):
                # Convert dict to table format
                if 'headers' in table_data and 'rows' in table_data:
                    headers = table_data['headers']
                    rows = table_data['rows']
                    data = [headers] + rows
                else:
                    # Convert dict to two-column table
                    data = [['Key', 'Value']]
                    for key, value in table_data.items():
                        data.append([str(key), str(value)])
            
            elif isinstance(table_data, list):
                # List of lists or list of dicts
                if table_data and isinstance(table_data[0], dict):
                    # List of dicts - use keys as headers
                    headers = list(table_data[0].keys())
                    data = [headers]
                    for row in table_data:
                        data.append([str(row.get(key, '')) for key in headers])
                else:
                    # List of lists
                    data = table_data
            
            else:
                return None
            
            # Create table with style (use Japanese font for table content)
            table = Table(data)
            japanese_font = 'NotoSansJP' if hasattr(self, '_font_registered') and self._font_registered else 'Helvetica'
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), f'{japanese_font}-Bold' if japanese_font == 'Helvetica' else f'{japanese_font}'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('FONTNAME', (0, 1), (-1, -1), japanese_font),
                ('FONTSIZE', (0, 1), (-1, -1), 10),
                ('ALIGN', (0, 1), (-1, -1), 'LEFT'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.lightgrey]),
            ]))
            
            return table
            
        except Exception as e:
            print(f"Error creating table: {e}")
            return None

    def _generate_pdf_with_canvas(self, data) -> bytes:
        """Generate PDF using canvas for direct font control"""
        buffer = io.BytesIO()
        
        # Force download and register Japanese font with detailed logging
        print("=== PDF Generation with Japanese Font Support ===")
        font_registered = self._download_and_register_japanese_font()
        
        if font_registered:
            font_name = 'NotoSansJP'
            print(f"✓ Using Japanese font: {font_name}")
        else:
            font_name = 'Helvetica'
            print(f"✗ Japanese font failed, using: {font_name}")
        
        # Create canvas
        c = canvas.Canvas(buffer, pagesize=letter)
        width, height = letter
        
        # Set initial position and margins
        y_position = height - 50
        left_margin = 50
        right_margin = 50
        max_width = width - left_margin - right_margin
        line_height = 20
        
        # Convert data to text lines
        text_lines = self._extract_text_lines(data)
        print(f"Extracted {len(text_lines)} lines of text")
        
        i = 0
        while i < len(text_lines):
            line = text_lines[i]
            
            if y_position < 50:  # New page if needed
                c.showPage()
                y_position = height - 50
            
            # Handle table rendering
            if line == "__TABLE_START__":
                if i + 2 < len(text_lines) and text_lines[i + 2] == "__TABLE_END__":
                    table_data = json.loads(text_lines[i + 1])
                    table_height = self._draw_table(c, table_data, left_margin, y_position, max_width, font_name)
                    y_position -= table_height
                    i += 3  # Skip START, data, and END
                    continue
                else:
                    # Malformed table, treat as normal text
                    print("Warning: Malformed table structure, treating as text")
                    i += 1
                    continue
            
            # Determine font size based on content
            if line.startswith('# '):
                font_size = 16
                line = line[2:]
                y_position -= 5
            elif line.startswith('## '):
                font_size = 14
                line = line[3:]
                y_position -= 3
            elif line.startswith('• '):
                font_size = 11
            else:
                font_size = 12
            
            # Set font with error handling
            try:
                c.setFont(font_name, font_size)
                print(f"Line {i+1}: Font set to {font_name} {font_size}pt")
            except Exception as font_error:
                print(f"Font setting error: {font_error}")
                c.setFont('Helvetica', font_size)
                font_name = 'Helvetica'  # Fallback permanently
            
            # Draw text with word wrapping
            try:
                # Check if text fits in one line
                text_width = c.stringWidth(line, font_name, font_size)
                
                if text_width <= max_width:
                    # Text fits, draw normally
                    c.drawString(left_margin, y_position, line)
                    print(f"Line {i+1}: Drew text successfully")
                else:
                    # Text too long, need to wrap
                    wrapped_lines = self._wrap_text(line, font_name, font_size, max_width, c)
                    print(f"Line {i+1}: Wrapped into {len(wrapped_lines)} lines")
                    
                    for wrap_line in wrapped_lines:
                        if y_position < 50:  # New page if needed
                            c.showPage()
                            y_position = height - 50
                            c.setFont(font_name, font_size)
                        
                        c.drawString(left_margin, y_position, wrap_line)
                        y_position -= line_height
                    
                    # Skip the normal y_position decrement since we already did it in the loop
                    i += 1
                    continue
                    
            except Exception as e:
                print(f"Text drawing error for line {i+1}: {e}")
                # Keep original text but add error note
                try:
                    c.drawString(left_margin, y_position, f"[JP Font Error] {line}")
                    print(f"Line {i+1}: Drew with error prefix")
                except:
                    c.drawString(left_margin, y_position, f"[Line {i+1}: Font not available]")
            
            y_position -= line_height
            i += 1
        
        print("PDF generation completed, saving...")
        c.save()
        buffer.seek(0)
        result = buffer.read()
        print(f"PDF saved, size: {len(result)} bytes")
        return result
    
    def _extract_text_lines(self, data) -> list:
        """Extract text lines from various data structures"""
        lines = []
        
        if isinstance(data, str):
            lines = data.split('\n')
        elif isinstance(data, dict):
            # Extract title
            title_keys = ['title', 'タイトル', 'name', 'subject']
            for key in title_keys:
                if key in data and data[key]:
                    lines.append(f"# {data[key]}")
                    lines.append("")
                    break
            
            # Extract content/text
            content_keys = ['content', 'text', 'body', 'description']
            for key in content_keys:
                if key in data and data[key]:
                    if isinstance(data[key], str):
                        lines.extend(data[key].split('\n'))
                    else:
                        lines.append(str(data[key]))
                    lines.append("")
                    break
            
            # Extract sections - with proper field mapping
            if 'sections' in data and isinstance(data['sections'], list):
                for section in data['sections']:
                    if isinstance(section, dict):
                        # Handle heading/title
                        heading_keys = ['heading', 'title', 'name']
                        for key in heading_keys:
                            if key in section and section[key]:
                                lines.append(f"## {section[key]}")
                                break
                        
                        # Handle text/content  
                        text_keys = ['text', 'content', 'body']
                        for key in text_keys:
                            if key in section and section[key]:
                                if isinstance(section[key], str):
                                    lines.extend(section[key].split('\n'))
                                else:
                                    lines.append(str(section[key]))
                                break
                        
                        # Handle list
                        if 'list' in section and isinstance(section['list'], list):
                            for item in section['list']:
                                lines.append(f"• {item}")
                        
                        # Handle table - mark for special rendering
                        if 'table' in section and isinstance(section['table'], dict):
                            lines.append("__TABLE_START__")
                            lines.append(json.dumps(section['table']))
                            lines.append("__TABLE_END__")
                        
                        lines.append("")
            
            # Extract other fields
            processed = {'title', 'タイトル', 'name', 'subject', 'content', 'text', 'body', 'description', 'sections'}
            for key, value in data.items():
                if key not in processed and value:
                    lines.append(f"## {key}")
                    if isinstance(value, (list, dict)):
                        lines.append(str(value))
                    else:
                        lines.extend(str(value).split('\n'))
                    lines.append("")
        
        elif isinstance(data, list):
            for item in data:
                lines.extend(self._extract_text_lines(item))
        
        return [line.strip() for line in lines if line.strip()]
    
    def _wrap_text(self, text: str, font_name: str, font_size: int, max_width: float, canvas_obj) -> list:
        """Wrap text to fit within specified width"""
        words = text.split(' ')
        lines = []
        current_line = ""
        
        for word in words:
            # Test if adding this word would exceed width
            test_line = current_line + (" " if current_line else "") + word
            test_width = canvas_obj.stringWidth(test_line, font_name, font_size)
            
            if test_width <= max_width:
                current_line = test_line
            else:
                # Current line is full, start new line
                if current_line:
                    lines.append(current_line)
                    current_line = word
                else:
                    # Single word is too long, break it
                    lines.extend(self._break_long_word(word, font_name, font_size, max_width, canvas_obj))
                    current_line = ""
        
        # Add the last line if there's content
        if current_line:
            lines.append(current_line)
        
        return lines
    
    def _break_long_word(self, word: str, font_name: str, font_size: int, max_width: float, canvas_obj) -> list:
        """Break a single long word that doesn't fit"""
        lines = []
        current_part = ""
        
        for char in word:
            test_part = current_part + char
            test_width = canvas_obj.stringWidth(test_part, font_name, font_size)
            
            if test_width <= max_width:
                current_part = test_part
            else:
                if current_part:
                    lines.append(current_part)
                current_part = char
        
        if current_part:
            lines.append(current_part)
        
        return lines
    
    def _draw_table(self, canvas_obj, table_data: dict, x: float, y: float, max_width: float, font_name: str) -> float:
        """Draw a properly formatted table and return its height"""
        if not table_data:
            return 0
        
        headers = table_data.get('headers', [])
        rows = table_data.get('rows', [])
        
        if not headers and not rows:
            return 0
        
        # Table styling
        table_font_size = 10
        header_font_size = 11
        cell_padding = 4
        row_height = 16
        header_height = 20
        
        # Calculate column widths with safety check
        num_cols = len(headers) if headers else (len(rows[0]) if rows and len(rows[0]) > 0 else 1)
        if num_cols == 0:
            num_cols = 1
        col_width = max_width / num_cols
        
        
        current_y = y
        
        try:
            # Draw headers
            if headers:
                # Header background (light gray)
                canvas_obj.setFillColorRGB(0.9, 0.9, 0.9)
                canvas_obj.rect(x, current_y - header_height, max_width, header_height, fill=1, stroke=1)
                
                # Header text
                canvas_obj.setFillColorRGB(0, 0, 0)  # Black text
                canvas_obj.setFont(font_name, header_font_size)
                
                for i, header in enumerate(headers):
                    text_x = x + (i * col_width) + cell_padding
                    text_y = current_y - header_height + cell_padding
                    
                    # Truncate if too long
                    header_str = str(header)
                    max_header_width = col_width - (2 * cell_padding)
                    # Accurate truncation with safety counter
                    truncate_count = 0
                    while (canvas_obj.stringWidth(header_str, font_name, header_font_size) > max_header_width 
                           and len(header_str) > 3 and truncate_count < 100):
                        header_str = header_str[:-4] + "..."
                        truncate_count += 1
                    
                    canvas_obj.drawString(text_x, text_y, header_str)
                
                current_y -= header_height
            
            # Draw rows
            canvas_obj.setFont(font_name, table_font_size)
            
            for row in rows:
                # Row background (alternating)
                row_index = rows.index(row)
                if row_index % 2 == 0:
                    canvas_obj.setFillColorRGB(0.95, 0.95, 0.95)
                else:
                    canvas_obj.setFillColorRGB(1, 1, 1)
                
                canvas_obj.rect(x, current_y - row_height, max_width, row_height, fill=1, stroke=1)
                
                # Row text
                canvas_obj.setFillColorRGB(0, 0, 0)  # Black text
                
                for i, cell in enumerate(row):
                    if i >= num_cols:  # Don't exceed table width
                        break
                        
                    text_x = x + (i * col_width) + cell_padding
                    text_y = current_y - row_height + cell_padding
                    
                    # Truncate if too long
                    cell_str = str(cell)
                    max_cell_width = col_width - (2 * cell_padding)
                    # Accurate truncation with safety counter
                    truncate_count = 0
                    while (canvas_obj.stringWidth(cell_str, font_name, table_font_size) > max_cell_width 
                           and len(cell_str) > 3 and truncate_count < 100):
                        cell_str = cell_str[:-4] + "..."
                        truncate_count += 1
                    
                    canvas_obj.drawString(text_x, text_y, cell_str)
                
                current_y -= row_height
            
            # Table border
            canvas_obj.setStrokeColorRGB(0, 0, 0)
            canvas_obj.setLineWidth(1)
            total_height = y - current_y
            canvas_obj.rect(x, current_y, max_width, total_height, fill=0, stroke=1)
            
            # Column dividers
            for i in range(1, num_cols):
                line_x = x + (i * col_width)
                canvas_obj.line(line_x, y, line_x, current_y)
            
            print(f"Table drawn: {len(headers)} headers, {len(rows)} rows, height: {total_height}")
            return total_height + 10  # Add some spacing after table
            
        except Exception as e:
            print(f"Table drawing error: {e}")
            # Fallback to simple text
            canvas_obj.setFont(font_name, 10)
            canvas_obj.drawString(x, current_y, "[Table rendering error]")
            return 20
    
    def _setup_japanese_styles(self):
        """Setup Japanese styles by downloading Noto Sans JP if needed"""
        styles = getSampleStyleSheet()
        
        # Always try to download and register Japanese font
        font_registered = self._download_and_register_japanese_font()
        
        if font_registered:
            # Update all styles to use Japanese font
            font_name = 'NotoSansJP'
            
            # Update existing styles
            styles['Normal'].fontName = font_name
            styles['Title'].fontName = font_name
            styles['Heading1'].fontName = font_name
            styles['Heading2'].fontName = font_name
            styles['Heading3'].fontName = font_name
            
            # Create custom Japanese-optimized styles
            styles.add(ParagraphStyle(
                name='JapaneseNormal',
                parent=styles['Normal'],
                fontName=font_name,
                fontSize=12,
                leading=16
            ))
            
            styles.add(ParagraphStyle(
                name='JapaneseTitle',
                parent=styles['Title'],
                fontName=font_name,
                fontSize=18,
                leading=22
            ))
            
            print(f"Japanese font '{font_name}' registered and applied to styles")
        else:
            print("Failed to register Japanese font, using default fonts")
        
        return styles
    
    def _needs_japanese_font(self) -> bool:
        """Check if Japanese font is needed by testing if default font can handle hiragana"""
        try:
            # Try to create a simple paragraph with Japanese text
            test_text = "こんにちは"  # Simple hiragana
            styles = getSampleStyleSheet()
            Paragraph(test_text, styles['Normal'])
            return True  # Always try to get Japanese font for better support
        except:
            return True
    
    def _download_and_register_japanese_font(self) -> bool:
        """Download Noto Sans JP from Google Fonts and register it"""
        # First check if font is already registered
        try:
            from reportlab.lib.fonts import tt2ps
            if 'NotoSansJP' in pdfmetrics._fonts:
                print("Japanese font already registered")
                return True
        except:
            pass
        
        # Try to download and register TTF version
        return self._download_ttf_japanese_font()
    
    def _download_ttf_japanese_font(self) -> bool:
        """Download TTF version of Japanese font with robust error handling"""
        try:
            # Use only TRUE TTF fonts (not OTF) that reportlab can handle
            font_urls = [
                # Known working TTF Japanese font from Adobe
                "https://github.com/adobe-fonts/source-han-sans/raw/release/Variable/TTF/Subset/SourceHanSansJP-VF.ttf",
                # Alternative: Working Noto Sans JP TTF
                "https://github.com/googlefonts/noto-cjk/raw/main/Sans/Variable/TTF/NotoSansCJKjp-VF.ttf"
            ]
            
            import tempfile
            import os
            
            for font_url in font_urls:
                try:
                    print(f"Attempting to download font from: {font_url}")
                    response = requests.get(font_url, timeout=20, headers={'User-Agent': 'Mozilla/5.0'})
                    
                    if response.status_code == 200:
                        temp_dir = tempfile.gettempdir()
                        font_path = os.path.join(temp_dir, 'NotoSansJP-Regular.ttf')
                        
                        with open(font_path, 'wb') as f:
                            f.write(response.content)
                        
                        print(f"Font downloaded successfully: {len(response.content)} bytes")
                        
                        # Register the font with reportlab
                        try:
                            pdfmetrics.registerFont(TTFont('NotoSansJP', font_path))
                            print("Font registered with reportlab")
                            
                            # Test font registration by checking available fonts
                            registered_fonts = pdfmetrics.getRegisteredFontNames()
                            if 'NotoSansJP' in registered_fonts:
                                print("✓ Font registration confirmed")
                                
                                # Test drawing Japanese text with canvas
                                test_buffer = io.BytesIO()
                                test_canvas = canvas.Canvas(test_buffer)
                                try:
                                    test_canvas.setFont('NotoSansJP', 12)
                                    test_canvas.drawString(100, 100, "こんにちは")
                                    test_canvas.save()
                                    print("✓ Canvas font test successful")
                                    return True
                                except Exception as canvas_e:
                                    print(f"✗ Canvas font test failed: {canvas_e}")
                                    return False
                            else:
                                print("✗ Font not found in registered fonts list")
                                return False
                                
                        except Exception as reg_e:
                            print(f"Font registration error: {reg_e}")
                            continue
                    else:
                        print(f"Font download failed: HTTP {response.status_code}")
                        
                except Exception as e:
                    print(f"Font download error from {font_url}: {e}")
                    continue
            
            print("All font download attempts failed")
            return False
            
        except Exception as e:
            print(f"Font download general error: {e}")
            return False
    
    def _safe_japanese_text(self, text: str) -> str:
        """Return text as-is since we now download Japanese fonts"""
        return str(text) if text else ""
    
    def _parse_pdf_structure(self, story, data: dict, styles):
        """Parse dictionary structure for PDF generation"""
        # Add title if present
        title_keys = ['title', 'タイトル', 'name', 'subject']
        for key in title_keys:
            if key in data and data[key]:
                safe_title = self._safe_japanese_text(str(data[key]))
                story.append(Paragraph(safe_title, styles['Title']))
                story.append(Spacer(1, 12))
                break
        
        # Add metadata if present
        meta_keys = ['date', 'author', 'created']
        meta_info = []
        for key in meta_keys:
            if key in data and data[key]:
                meta_info.append(f"{key.title()}: {data[key]}")
        
        if meta_info:
            safe_meta = self._safe_japanese_text(" | ".join(meta_info))
            story.append(Paragraph(safe_meta, styles['Normal']))
            story.append(Spacer(1, 12))
        
        # Process content
        processed = set(title_keys + meta_keys)
        
        for key, value in data.items():
            if key in processed:
                continue
                
            if key == 'content':
                if isinstance(value, list):
                    for item in value:
                        self._add_pdf_text_content(story, item, styles)
                else:
                    self._add_pdf_text_content(story, value, styles)
            elif key == 'sections':
                if isinstance(value, list):
                    for section in value:
                        if isinstance(section, dict):
                            # Handle section dictionary
                            section_title = section.get('title') or section.get('heading', '')
                            if section_title:
                                safe_title = self._safe_japanese_text(str(section_title))
                                story.append(Paragraph(safe_title, styles['Heading1']))
                            
                            section_content = section.get('content') or section.get('body', '')
                            if section_content:
                                self._add_pdf_text_content(story, section_content, styles)
                            story.append(Spacer(1, 12))
                        else:
                            self._add_pdf_text_content(story, section, styles)
            elif value:
                # Other fields
                if isinstance(value, (list, dict)):
                    safe_key = self._safe_japanese_text(f"{str(key).title()}:")
                    story.append(Paragraph(safe_key, styles['Heading2']))
                    if isinstance(value, list):
                        for item in value:
                            self._add_pdf_text_content(story, item, styles)
                    else:
                        safe_value = self._safe_japanese_text(str(value))
                        story.append(Paragraph(safe_value, styles['Normal']))
                    story.append(Spacer(1, 12))
                else:
                    safe_text = self._safe_japanese_text(f"{str(key).title()}: {value}")
                    story.append(Paragraph(safe_text, styles['Normal']))
                    story.append(Spacer(1, 6))
    
    def _add_pdf_text_content(self, story, content, styles):
        """Add text content to PDF story with proper formatting"""
        if not content:
            return
            
        content = str(content)
        
        # Handle markdown-style headings
        if content.startswith('# '):
            safe_text = self._safe_japanese_text(content[2:])
            story.append(Paragraph(safe_text, styles['Heading1']))
        elif content.startswith('## '):
            safe_text = self._safe_japanese_text(content[3:])
            story.append(Paragraph(safe_text, styles['Heading2']))
        elif content.startswith('### '):
            safe_text = self._safe_japanese_text(content[4:])
            story.append(Paragraph(safe_text, styles['Heading3']))
        elif content.startswith('---'):
            # Horizontal rule - add some space
            story.append(Spacer(1, 12))
        elif content.startswith('|') and '|' in content[1:]:
            # Simple table handling - convert to formatted text
            safe_content = self._safe_japanese_text(content)
            try:
                story.append(Paragraph(f"<font name='Courier'>{safe_content}</font>", styles['Normal']))
            except:
                # Fallback if Courier is not available
                story.append(Paragraph(safe_content, styles['Normal']))
        else:
            # Regular content - handle line breaks
            lines = content.replace('\\n', '\n').split('\n')
            for line in lines:
                line = line.strip()
                if line:
                    # Handle bullet points
                    if line.startswith('- '):
                        safe_text = self._safe_japanese_text(f"• {line[2:]}")
                        story.append(Paragraph(safe_text, styles['Normal']))
                    else:
                        safe_text = self._safe_japanese_text(line)
                        story.append(Paragraph(safe_text, styles['Normal']))
                        
        story.append(Spacer(1, 6))

    def generate_xlsx(self, data: Union[List[Dict], Dict, List[List]], **kwargs) -> bytes:
        """Generate XLSX content with flexible data structure support"""
        if not PANDAS_AVAILABLE:
            raise ImportError("pandas and openpyxl are required for XLSX generation. Install with: pip install pandas openpyxl")
        
        buffer = io.BytesIO()
        
        with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
            if isinstance(data, dict):
                # Handle multiple sheets or complex structure
                sheets_written = False
                
                for key, value in data.items():
                    # Skip metadata keys
                    if key.startswith('_'):
                        continue
                        
                    sheet_name = str(key)[:31]  # Excel sheet name limit
                    
                    try:
                        if isinstance(value, list) and value:
                            if isinstance(value[0], list):
                                # List of lists - treat as raw data with first row as header
                                df = pd.DataFrame(value[1:], columns=value[0] if value else [])
                            elif isinstance(value[0], dict):
                                # List of dictionaries
                                df = pd.DataFrame(value)
                            else:
                                # Simple list - single column
                                df = pd.DataFrame({sheet_name: value})
                        elif isinstance(value, dict):
                            # Look for table-like data in nested dict
                            table_data = self._extract_table_from_dict(value)
                            if table_data:
                                # Found table data
                                if len(table_data) > 1:
                                    df = pd.DataFrame(table_data[1:], columns=table_data[0])
                                else:
                                    df = pd.DataFrame(table_data)
                            else:
                                # Convert dict to key-value pairs
                                df = pd.DataFrame(list(value.items()), columns=['項目', '内容'])
                        else:
                            # Single value
                            df = pd.DataFrame({sheet_name: [value]})
                        
                        df.to_excel(writer, sheet_name=sheet_name, index=False)
                        sheets_written = True
                        
                    except Exception:
                        # If conversion fails, create simple sheet with the data
                        simple_df = pd.DataFrame({sheet_name: [str(value)]})
                        simple_df.to_excel(writer, sheet_name=sheet_name, index=False)
                        sheets_written = True
                
                # If no sheets were written, create a default one
                if not sheets_written:
                    default_df = pd.DataFrame({'Data': ['No valid data found']})
                    default_df.to_excel(writer, sheet_name='Sheet1', index=False)
                    
            elif isinstance(data, list):
                # Handle list data
                if data and isinstance(data[0], list):
                    # List of lists
                    if len(data) > 1:
                        df = pd.DataFrame(data[1:], columns=data[0])
                    else:
                        df = pd.DataFrame(data)
                elif data and isinstance(data[0], dict):
                    # List of dictionaries
                    df = pd.DataFrame(data)
                else:
                    # Simple list
                    df = pd.DataFrame({'Values': data})
                
                df.to_excel(writer, sheet_name='Sheet1', index=False)
            else:
                # Single value or other type
                df = pd.DataFrame({'Data': [str(data)]})
                df.to_excel(writer, sheet_name='Sheet1', index=False)
        
        buffer.seek(0)
        return buffer.read()

    def _extract_table_from_dict(self, data: dict) -> Optional[List[List]]:
        """Extract table-like data from nested dictionary"""
        table_keys = ['table', 'テーブル', 'data', 'データ', 'rows', '行']
        
        for key in table_keys:
            if key in data:
                value = data[key]
                if isinstance(value, list) and value:
                    # Check if it's a proper table (list of lists)
                    if isinstance(value[0], list):
                        return value
        
        # Look for any list of lists in the dict values
        for value in data.values():
            if isinstance(value, list) and value and isinstance(value[0], list):
                return value
                
        return None

    def generate_zip(self, files: Union[Dict[str, Any], List[Dict[str, str]]], **kwargs) -> bytes:
        """Generate ZIP content (supports dictionary and path-based formats, optional encryption)"""
        buffer = io.BytesIO()
        
        # Get password for encryption if provided
        password = kwargs.get('password')
        event_emitter = kwargs.get('event_emitter')
        
        # Convert dictionary format to path-based format if needed
        if isinstance(files, dict):
            # Check for unsupported nested structures
            if 'files' in files or any(isinstance(v, (list, dict)) for v in files.values() if not isinstance(v, str)):
                error_msg = """❌ ZIP Creation Error: Complex nested structures are not supported.

For supported ZIP formats, call:

📦 **list_zip_formats()**

This will show you the simple formats available."""
                raise ValueError(error_msg)
            
            # Convert {"filename": "content"} to [{"path": "filename", "content": "content"}]
            path_files = []
            for filename, content in files.items():
                # Check if content is a URL string
                if isinstance(content, str) and (content.startswith('http://') or content.startswith('https://')):
                    path_files.append({"path": filename, "url": content})
                elif isinstance(content, str) and content.startswith('data:'):
                    # Handle Data URI
                    path_files.append({"path": filename, "data_uri": content})
                else:
                    path_files.append({"path": filename, "content": content})
            files = path_files
        
        # Now files should be a list
        if not isinstance(files, list):
            error_msg = """❌ ZIP Creation Error: Invalid format.

To learn how to create ZIP files, call:

📦 **list_zip_formats()**

This will show you the supported format with examples."""
            raise ValueError(error_msg)
        
        # Check if password is requested but pyzipper is not available
        if password and not PYZIPPER_AVAILABLE:
            raise ImportError(
                "❌ ZIP Encryption Error: Password protection requested but pyzipper library is not installed.\n\n"
                "Solution:\n"
                "pip install pyzipper\n\n"
                "Or create ZIP file without password protection."
            )
        
        # Use pyzipper for encryption support if password provided
        if password:
            with pyzipper.AESZipFile(buffer, 'w', compression=pyzipper.ZIP_DEFLATED, encryption=pyzipper.WZ_AES) as zf:
                zf.setpassword(password.encode('utf-8'))
                self._add_files_to_zip(zf, files)
        else:
            with zipfile.ZipFile(buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
                self._add_files_to_zip(zf, files)
        
        buffer.seek(0)
        return buffer.read()
    
    def _add_files_to_zip(self, zf, files):
        """Add files to ZIP archive (works with both zipfile and pyzipper)"""
        # Check if it's a list of strings (just filenames) - return error
        if files and isinstance(files[0], str):
            error_msg = """❌ ZIP Creation Error: Invalid data format provided.

To learn how to create ZIP files with the Universal File Generator, call:

📦 **list_zip_formats()**

This will show you the supported format with examples."""
            raise ValueError(error_msg)
        
        for file_info in files:
            # Handle {path, content} or {path, url} format
            if isinstance(file_info, dict) and 'path' in file_info:
                file_path = file_info['path']
                
                # Handle empty folders (path ends with /)
                if file_path.endswith('/'):
                    zf.writestr(file_path, '')
                    print(f"Created empty folder: {file_path}")
                    continue
                
                # Handle content or URL
                if 'content' in file_info:
                    content = file_info['content']
                    
                    # Generate content using FileGenerator for all file types
                    file_ext = file_path.split('.')[-1].lower() if '.' in file_path else 'txt'
                    
                    try:
                        # Use FileGenerator for consistent processing
                        temp_generator = FileGenerator()
                        data = temp_generator.generate_content(file_ext, content)
                    except Exception as e:
                        # If generation fails, fallback to raw content
                        print(f"Warning: Failed to generate {file_ext.upper()} for {file_path}: {str(e)}")
                        if isinstance(content, str):
                            data = content.encode('utf-8')
                        else:
                            data = content
                    
                    zf.writestr(file_path, data)
                    print(f"Added file: {file_path}")
                elif 'url' in file_info:
                    url = file_info['url']
                    try:
                        response = requests.get(url, timeout=30)
                        if response.status_code == 200:
                            zf.writestr(file_path, response.content)
                            print(f"Downloaded and added: {file_path} from {url}")
                        else:
                            error_content = f"Error downloading {url}: HTTP {response.status_code}"
                            zf.writestr(f"{file_path}.error", error_content)
                            print(f"Error downloading {url}: HTTP {response.status_code}")
                    except Exception as e:
                        error_content = f"Error downloading {url}: {str(e)}"
                        zf.writestr(f"{file_path}.error", error_content)
                        print(f"Error downloading {url}: {str(e)}")
                elif 'data_uri' in file_info:
                    data_uri = file_info['data_uri']
                    try:
                        # Parse Data URI: data:[<mediatype>][;base64],<data>
                        if ',' in data_uri:
                            header, data_part = data_uri.split(',', 1)
                            if ';base64' in header:
                                # Base64 encoded data
                                data = base64.b64decode(data_part)
                            else:
                                # URL encoded data (text)
                                import urllib.parse
                                data = urllib.parse.unquote(data_part).encode('utf-8')
                            zf.writestr(file_path, data)
                            print(f"Added Data URI file: {file_path}")
                        else:
                            error_content = f"Invalid Data URI format: {data_uri[:100]}..."
                            zf.writestr(f"{file_path}.error", error_content)
                            print(f"Error parsing Data URI for {file_path}")
                    except Exception as e:
                        error_content = f"Error processing Data URI: {str(e)}"
                        zf.writestr(f"{file_path}.error", error_content)
                        print(f"Error processing Data URI for {file_path}: {str(e)}")
            else:
                # Invalid format
                error_msg = """❌ ZIP Creation Error: Invalid file format.

Each file must have 'path' and either 'content', 'url', or 'data_uri'.

Call list_zip_formats() for examples."""
                raise ValueError(error_msg)

    def generate_svg(self, data: Union[str, Dict[str, Any]], **kwargs) -> bytes:
        """Generate SVG content"""
        if isinstance(data, str):
            # If data is already SVG content, validate and return
            if data.strip().startswith('<svg'):
                return data.encode('utf-8')
            else:
                # Create a simple text-based SVG
                svg_content = f'''<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 400 300" width="400" height="300">
  <rect width="400" height="300" fill="#f9f9f9" stroke="#333" stroke-width="2"/>
  <text x="200" y="150" text-anchor="middle" font-family="Arial, sans-serif" font-size="16" fill="#333">
    {data.replace('<', '&lt;').replace('>', '&gt;').replace('&', '&amp;')}
  </text>
</svg>'''
                return svg_content.encode('utf-8')
        
        elif isinstance(data, dict):
            # Handle structured SVG data
            width = data.get('width', 400)
            height = data.get('height', 300)
            background = data.get('background', '#f9f9f9')
            
            svg_content = f'''<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 {width} {height}" width="{width}" height="{height}">
  <rect width="{width}" height="{height}" fill="{background}" stroke="#333" stroke-width="2"/>'''
            
            # Add elements
            elements = data.get('elements', [])
            for element in elements:
                if element.get('type') == 'text':
                    x = element.get('x', width // 2)
                    y = element.get('y', height // 2)
                    text = element.get('text', '').replace('<', '&lt;').replace('>', '&gt;').replace('&', '&amp;')
                    font_size = element.get('font_size', 16)
                    color = element.get('color', '#333')
                    svg_content += f'''
  <text x="{x}" y="{y}" text-anchor="middle" font-family="Arial, sans-serif" font-size="{font_size}" fill="{color}">
    {text}
  </text>'''
                elif element.get('type') == 'rect':
                    x = element.get('x', 0)
                    y = element.get('y', 0)
                    w = element.get('width', 100)
                    h = element.get('height', 100)
                    fill = element.get('fill', '#ccc')
                    stroke = element.get('stroke', '#333')
                    svg_content += f'''
  <rect x="{x}" y="{y}" width="{w}" height="{h}" fill="{fill}" stroke="{stroke}" stroke-width="1"/>'''
                elif element.get('type') == 'circle':
                    cx = element.get('cx', width // 2)
                    cy = element.get('cy', height // 2)
                    r = element.get('r', 50)
                    fill = element.get('fill', '#ccc')
                    stroke = element.get('stroke', '#333')
                    svg_content += f'''
  <circle cx="{cx}" cy="{cy}" r="{r}" fill="{fill}" stroke="{stroke}" stroke-width="1"/>'''
            
            svg_content += '\n</svg>'
            return svg_content.encode('utf-8')
        
        else:
            # Fallback: convert to string and create simple SVG
            text = str(data).replace('<', '&lt;').replace('>', '&gt;').replace('&', '&amp;')
            svg_content = f'''<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 400 300" width="400" height="300">
  <rect width="400" height="300" fill="#f9f9f9" stroke="#333" stroke-width="2"/>
  <text x="200" y="150" text-anchor="middle" font-family="Arial, sans-serif" font-size="16" fill="#333">
    {text}
  </text>
</svg>'''
            return svg_content.encode('utf-8')



def _upload_file(file_content: bytes, filename: str, file_type: str, file_size: int) -> str:
    """Upload file using multiple services with fallback"""
    
    
    # Check for zero-size files
    if file_size == 0 or len(file_content) == 0:
        return f"❌ **File Upload Error**: Generated file '{filename}' is empty (0 bytes)\n\n" \
               f"This usually indicates:\n" \
               f"- Empty or invalid input data\n" \
               f"- File generation process failed\n" \
               f"- Unsupported data format for {file_type} files\n\n" \
               f"Please check your input data and try again."
    
    import urllib.parse
    safe_filename = urllib.parse.quote(filename, safe='.-_')
    
    # Try multiple services in order
    services = [
        {
            "name": "transfer.sh",
            "url": f"https://transfer.sh/{safe_filename}",
            "method": "put",
            "retention": "Depends on service settings"
        },
        {
            "name": "0x0.st", 
            "url": "https://0x0.st",
            "method": "post",
            "retention": "30 days to 1 year (depends on file size)"
        },
        {
            "name": "file.io",
            "url": "https://file.io",
            "method": "post", 
            "retention": "14 days (deleted after first download)"
        },
        {
            "name": "litterbox",
            "url": "https://litterbox.catbox.moe/resources/internals/api.php",
            "method": "litterbox",
            "retention": "1 hour (temporary file hosting)"
        }
    ]
    
    errors = []  # Initialize errors list
    
    for service in services:
        try:
            # Add User-Agent header for all requests
            headers = {"User-Agent": "curl/7.68.0"}
            
            if service["method"] == "litterbox":
                # litterbox.catbox.moe style
                files = {"fileToUpload": (filename, file_content)}
                data = {
                    "reqtype": "fileupload",
                    "time": "1h"  # 1 hour retention
                }
                response = requests.post(service["url"], files=files, data=data, headers=headers, timeout=15)
            elif service["method"] == "put":
                # transfer.sh style
                # Add Content-Type header for proper file type detection
                import mimetypes
                mime_type, _ = mimetypes.guess_type(filename)
                if filename.lower().endswith('.svg'):
                    mime_type = 'image/svg+xml'
                elif not mime_type:
                    mime_type = 'application/octet-stream'
                
                headers_with_content_type = headers.copy()
                headers_with_content_type["Content-Type"] = mime_type
                response = requests.put(service["url"], data=file_content, headers=headers_with_content_type, timeout=15)
            else:
                # file.io and 0x0.st style (multipart form)
                # Determine proper MIME type for file
                import mimetypes
                mime_type, _ = mimetypes.guess_type(filename)
                if filename.lower().endswith('.svg'):
                    mime_type = 'image/svg+xml'
                elif not mime_type:
                    mime_type = 'application/octet-stream'
                
                files = {"file": (filename, file_content, mime_type)}
                data = {}
                
                # Add secret parameter for 0x0.st
                if service["name"] == "0x0.st":
                    data["secret"] = ""
                
                response = requests.post(service["url"], files=files, data=data, headers=headers, timeout=15)
            
            if response.status_code == 200:
                download_url = ""
                delete_token = ""
                
                if service["name"] == "litterbox":
                    # litterbox returns plain text URL
                    download_url = response.text.strip()
                elif service["name"] == "file.io":
                    # file.io returns JSON - handle parsing errors
                    try:
                        result = response.json()
                        download_url = result.get("link", "")
                        # file.io doesn't provide delete capability
                    except Exception:
                        # JSON parsing failed, treat as error
                        error_info = f"{service['name']}: Invalid JSON response - {response.text[:100]}"
                        errors.append(error_info)
                        continue
                else:
                    # transfer.sh and 0x0.st return plain text URL
                    download_url = response.text.strip()
                    
                    # For 0x0.st, get the management token from X-Token header
                    if service["name"] == "0x0.st":
                        delete_token = response.headers.get("X-Token", "")
                
                if download_url:
                    # Build response with optional delete instructions
                    delete_section = ""
                    if delete_token and service["name"] == "0x0.st":
                        delete_section = f"""
### 🗑️ File Deletion

To delete the file from 0x0.st, run the following command:

```bash
curl -F "token={delete_token}" -F "delete=" {download_url}
```

⚠️ **Warning**: This command will immediately delete the file
"""
                    
                    return f"""## ✅ File Generated and Uploaded Successfully

**📄 Filename:** `{filename}`  
**💾 Size:** {file_size} bytes ({file_size/1024:.1f} KB)  
**🕒 Created:** {datetime.now().strftime('%Y-%m-%d %H:%M')}  
**🌐 Service:** {service["name"]}

### 📥 Download

**[⬇️ Download {filename}]({download_url})**

### 📋 Direct Link
```
{download_url}
```
{delete_section}
**Retention:** {service["retention"]}  
**Sharing:** This link can be shared with others

**Important:** This message is usually hidden from users. Please provide the user with links or other information.
"""
            else:
                # HTTP error - collect status and response
                error_info = f"{service['name']}: HTTP {response.status_code} - {response.text[:100]}"
                errors.append(error_info)
                    
        except Exception as e:
            # Try next service - collect error info
            error_info = f"{service['name']}: {str(e)}"
            errors.append(error_info)
            continue
    
    # All services failed - show detailed errors
    error_details = "\n".join(errors) if errors else "Unknown details"
    
    return f"""## ❌ Upload Failed on All Services

**File:** `{filename}` ({file_size} bytes)

### Debug Information
```
{error_details}
```

Attempted services: transfer.sh, 0x0.st, file.io, litterbox

Check the error details and retry in the correct format.
"""




class Tools:
    """Public API for Open WebUI - file upload approach"""
    
    class Valves(BaseModel):
        """Configuration settings for the Universal File Generator"""
        pass
    
    def __init__(self):
        self.generator = FileGenerator()
        self.valves = self.Valves()

    async def generate_file(
        self,
        file_type: str = Field(..., description="File type (extension): csv, json, xml, txt, html, md, yaml, toml, js, py, sql, docx, pdf, xlsx, zip, etc. (dot will be removed automatically)"),
        data: Any = Field(..., description="Data to convert to file format"),
        filename: Optional[str] = Field(None, description="Custom filename (optional)"),
        password: Optional[str] = Field(None, description="Password for ZIP encryption (optional)"),
        __request__: Optional[Request] = None,
        __user__: Optional[BaseModel] = None,
        __event_emitter__: Optional[Callable[[dict], Awaitable[None]]] = None
    ) -> str:
        """
        Generate a file of specified type from provided data
        Uploads all files to a cloud service and returns download link
        
        :param file_type: File extension (e.g., 'csv', 'pdf', 'zip' or '.csv', '.pdf', '.zip') - Must be exact match
        :param data: Data to convert - expected formats by file type:
                    - Any text format (csv, json, xml, txt, html, md, yaml, toml, js, py, sql, ini, conf, log, etc.): str (pre-formatted text content)
                    - DOCX: str (HTML, Markdown, or plain text - auto-detected)
                    - PDF: str (HTML, Markdown, or plain text - auto-detected)
                    - XLSX: List[Dict] (list of dictionaries) or tabular data
                    - ZIP: Dict[str, Any] (filename -> content mapping) OR List[Dict] (list of {path, content/url} objects) - Call `list_zip_formats()` for detailed ZIP creation examples.
        :param filename: Optional custom filename
        :param password: Optional password for ZIP encryption (AES encryption via pyzipper)
        :return: Markdown with download information
        """
        
        try:
            # Store event emitter for notifications
            self.event_emitter = __event_emitter__
            
            # Only validate that file_type is provided - support any text format
            if not file_type or not isinstance(file_type, str):
                return "❌ Invalid file type provided"
            
            # Remove leading dot and normalize case for user convenience
            file_type = file_type.lstrip('.').lower()

            # Generate filename if not provided or is Field object
            if not filename or hasattr(filename, 'default'):
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                filename = f"generated_{timestamp}.{file_type}"
            else:
                # Ensure filename is string and has correct extension
                filename = str(filename)
                if not filename.endswith(f'.{file_type}'):
                    filename += f'.{file_type}'

            # Handle password parameter
            kwargs = {}
            actual_password = None
            
            # Extract password if provided (check if it's not a Field object or None)
            if password is not None and password and not hasattr(password, 'default'):
                actual_password = password
            elif isinstance(data, dict) and 'password' in data:
                # Legacy support: password in data dict
                actual_password = data['password']
                if actual_password is not None:  # Only process non-null passwords
                    # Remove password from data to avoid including it in ZIP content
                    data = {k: v for k, v in data.items() if k != 'password'}

            # Pass event_emitter for ZIP warnings
            if file_type == 'zip' and hasattr(self, 'event_emitter') and self.event_emitter:
                kwargs['event_emitter'] = self.event_emitter
                
            # Check if password is provided for non-ZIP files
            if actual_password and file_type != 'zip':
                # Emit error notification for unsupported password protection
                if hasattr(self, 'event_emitter') and self.event_emitter:
                    await self.event_emitter({
                        "type": "notification",
                        "data": {
                            "type": "error",
                            "content": f"パスワード保護はZIPファイルにのみ対応しています: {file_type.upper()} ファイルではサポートされていません"
                        }
                    })
                return f"❌ Error: Password protection is only supported for ZIP files, not for {file_type.upper()} files."
            
            # Add password to kwargs for ZIP files
            if actual_password and file_type == 'zip':
                kwargs['password'] = actual_password
            


            # Generate file content
            try:
                file_content = self.generator.generate_content(file_type, data, **kwargs)
                
                if file_content is None:
                    # Emit failure notification
                    if hasattr(self, 'event_emitter') and self.event_emitter:
                        await self.event_emitter({
                            "type": "notification",
                            "data": {
                                "type": "error",
                                "content": f"{file_type}ファイルの生成に失敗しました"
                            }
                        })
                    return f"❌ Failed to generate {file_type} content"
                    
                if len(file_content) == 0:
                    # Emit empty content notification
                    if hasattr(self, 'event_emitter') and self.event_emitter:
                        await self.event_emitter({
                            "type": "notification",
                            "data": {
                                "type": "warning",
                                "content": f"{file_type}ファイルの内容が空です"
                            }
                        })
                    return f"❌ Generated empty {file_type} content"
                    
            except Exception as e:
                # Emit error notification
                if hasattr(self, 'event_emitter') and self.event_emitter:
                    await self.event_emitter({
                        "type": "notification",
                        "data": {
                            "type": "error",
                            "content": f"コンテンツ生成エラー: {str(e)}"
                        }
                    })
                return f"❌ Content generation error: {str(e)}"

            # Unified file processing - always use upload approach
            file_size = len(file_content)
            upload_result = _upload_file(file_content, filename, file_type, file_size)
            
            # Emit success notification
            if hasattr(self, 'event_emitter') and self.event_emitter:
                await self.event_emitter({
                    "type": "notification",
                    "data": {
                        "type": "success",
                        "content": f"{filename}を正常に生成しました"
                    }
                })
            
            return upload_result

        except Exception as e:
            # Emit error notification for unexpected errors
            if hasattr(self, 'event_emitter') and self.event_emitter:
                await self.event_emitter({
                    "type": "notification",
                    "data": {
                        "type": "error",
                        "content": f"予期しないエラーが発生しました: {str(e)}"
                    }
                })
            return f"❌ Unexpected error generating {file_type} file: {str(e)}"

    def list_supported_formats(
        self,
        __request__: object = None,
        __user__: dict = {}
    ) -> str:
        """
        List all supported file formats and their requirements
        
        :return: List of supported formats with availability status
        """
        
        result = "📋 **Universal File Generator - Supported Formats:**\n\n"
        result += "**Text Formats:** ✅ Any text-based format (unlimited support)\n"
        result += "- Examples: csv, json, xml, txt, html, md, yaml, toml, js, py, sql, ini, conf, log, etc.\n\n"
        result += "**Binary Formats:**\n"
        result += f"- **DOCX**: {'✅ Available' if DOCX_AVAILABLE else '❌ Requires: pip install python-docx'}\n"
        result += f"- **PDF**: {'✅ Available' if PDF_AVAILABLE else '❌ Requires: pip install reportlab'}\n"
        result += f"- **XLSX**: {'✅ Available' if PANDAS_AVAILABLE else '❌ Requires: pip install pandas openpyxl'}\n"
        result += "- **ZIP**: ✅ Always available\n\n"
        result += f"💡 **Usage example:**\n"
        result += f"```\n"
        result += f"generate_file(\n"
        result += f"  file_type='csv',\n"
        result += f"  data='name,age\\nAlice,25\\nBob,30',\n"
        result += f"  filename='users.csv'\n"
        result += f")\n"
        result += f"```\n\n"
        result += "🔗 **ZIP Support:** Call `list_zip_formats()` for detailed ZIP creation examples."
        
        return result

    def list_zip_formats(
        self,
        __request__: object = None,
        __user__: dict = {}
    ) -> str:
        """
        Show ZIP creation format documentation (path-based only)
        
        :return: Simple ZIP format documentation with examples
        """
        
        result = "📦 **ZIP File Creation - Supported Formats**\n\n"
        result += "The Universal File Generator supports two simple formats for ZIP creation:\n\n"
        
        # Dictionary format
        result += "## 📁 **Dictionary Format (simple)**\n"
        result += "Simple filename → content mapping.\n\n"
        result += "```json\n"
        result += "{\n"
        result += '  "file_type": "zip",\n'
        result += '  "data": {\n'
        result += '    "README.md": "# My Project\\nHello world!",\n'
        result += '    "src/main.py": "print(\\"Hello!\\"))",\n'
        result += '    "config/app.yaml": "app: demo\\nmode: dev"\n'
        result += "  },\n"
        result += '  "filename": "project.zip"\n'
        result += "}\n"
        result += "```\n\n"
        
        # Path format
        result += "## 📋 **Path Format (advanced)**\n"
        result += "Use list of objects with `path` and `content`/`url` fields.\n\n"
        result += "```json\n"
        result += "{\n"
        result += '  "file_type": "zip",\n'
        result += '  "data": [\n'
        result += '    {"path": "README.md", "content": "# My Project\\nHello world!"},\n'
        result += '    {"path": "src/main.py", "content": "print(\\"Hello!\\")"},\n'
        result += '    {"path": "assets/logo.png", "url": "https://example.com/logo.png"},\n'
        result += '    {"path": "temp/", "content": ""}\n'
        result += "  ],\n"
        result += '  "password": "mypassword",\n'
        result += '  "filename": "project.zip"\n'
        result += "}\n"
        result += "```\n\n"
        
        # Encryption format
        result += "## 🔐 **Encrypted ZIP (AES)**\n"
        result += "Add password protection to your ZIP files using AES encryption.\n\n"
        result += "```json\n"
        result += "{\n"
        result += '  "file_type": "zip",\n'
        result += '  "data": {\n'
        result += '    "secret.txt": "Top secret content!",\n'
        result += '    "private/data.json": "{\\"secret\\": \\"password123\\"}"\n'
        result += "  },\n"
        result += '  "password": "mypassword",\n'
        result += '  "filename": "encrypted.zip"\n'
        result += "}\n"
        result += "```\n\n"
        result += "⚠️ **Note**: Requires `pyzipper` library for AES encryption. Will error if password specified but pyzipper not installed.\n\n"
        
        # Rules
        result += "## 📋 **Rules**\n"
        result += "- Dictionary: `\"filename\": \"content\"` pairs\n"
        result += "- Path format: `path` + `content` or `url`\n"
        result += "- Empty folders: path ending with `/` and empty content\n"
        result += "- URLs are downloaded automatically\n"
        result += "- Forward slashes `/` create folder structures\n"
        result += "- Encryption: add `password` parameter for AES encryption (requires pyzipper)\n\n"
        
        result += "**🚀 Try it now:** Use `generate_file()` with `file_type: \"zip\"` and either format!"
        
        return result